import os
import re
import unicodedata
import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import webbrowser
import time

# Import các thư viện Selenium và thư viện chống phát hiện
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException
import undetected_chromedriver as uc # <<< SỬ DỤNG THƯ VIỆN MỚI

TEMPLATE_PATH = "mau_ct01.docx"

def normalize_filename(text):
    """Chuẩn hóa chuỗi thành tên file an toàn."""
    nfkd = unicodedata.normalize('NFKD', text)
    no_diacritics = "".join([c for c in nfkd if not unicodedata.combining(c)])
    safe = re.sub(r"[^A-Za-z0-9]+", "_", no_diacritics)
    return safe.strip("_")

def set_cell_text(cell, text, size=16, is_digit=False):
    """Điền văn bản vào một ô trong bảng Word với định dạng cụ thể."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if is_digit:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def fill_template(main_info, members):
    """Điền thông tin từ giao diện vào file mẫu Word CT01."""
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(f"Không tìm thấy file mẫu: {TEMPLATE_PATH}")
    doc = Document(TEMPLATE_PATH)
    tables = doc.tables
    def add_formatted_run(paragraph, text_to_add):
        run = paragraph.add_run(text_to_add)
        run.font.size = Pt(16)
        run.font.name = "Times New Roman"
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt.startswith("1. Họ, chữ đệm và tên"):
            p.text = "1. Họ, chữ đệm và tên: "
            add_formatted_run(p, main_info.get('name','').upper())
        elif "2. Ngày, tháng, năm sinh" in txt and "3. Giới tính" in txt:
            p.text = "2. Ngày, tháng, năm sinh: "
            add_formatted_run(p, main_info.get('dob',''))
            p.add_run("       3. Giới tính: ")
            add_formatted_run(p, main_info.get('gender',''))
        elif txt.startswith("5. Số điện thoại"):
            p.text = "5. Số điện thoại liên hệ: "
            add_formatted_run(p, main_info.get('phone',''))
            p.add_run("        6. Email:")
        elif txt.startswith("7. Họ, chữ đệm và tên chủ hộ"):
            p.text = "7. Họ, chữ đệm và tên chủ hộ: "
            add_formatted_run(p, main_info.get('owner_name',''))
            p.add_run("   8. Mối quan hệ với chủ hộ: ")
            add_formatted_run(p, main_info.get('owner_relation',''))

    cccd_person = main_info.get("id", "")
    if len(tables) > 0:
        table4 = tables[0]
        for i in range(12):
            cell_index = i + 1
            if cell_index < len(table4.rows[0].cells):
                if i < len(cccd_person): set_cell_text(table4.cell(0, cell_index), str(cccd_person[i]), is_digit=True)
                else: set_cell_text(table4.cell(0, cell_index), "", is_digit=True)

    cccd_owner = main_info.get("owner_id","")
    if len(tables) > 1:
        table9 = tables[1]
        for i in range(12):
            cell_index = i + 1
            if cell_index < len(table9.rows[0].cells):
                if i < len(cccd_owner): set_cell_text(table9.cell(0, cell_index), str(cccd_owner[i]), is_digit=True)
                else: set_cell_text(table9.cell(0, cell_index), "", is_digit=True)

    if len(tables) > 2:
        table11 = tables[2]
        for row_idx, member in enumerate(members, start=1):
            if row_idx < len(table11.rows):
                cells = table11.rows[row_idx].cells
                set_cell_text(cells[0], str(row_idx))
                set_cell_text(cells[1], member.get("name",""))
                set_cell_text(cells[2], member.get("dob",""))
                set_cell_text(cells[3], member.get("gender",""))
                set_cell_text(cells[4], member.get("id_number",""))
                set_cell_text(cells[5], member.get("relation",""))

    base_name = normalize_filename(main_info.get("name","Nguoi_dan"))
    counter = 1
    out_docx = f"CT01-{base_name}-{counter}.docx"
    while os.path.exists(out_docx):
        counter += 1
        out_docx = f"CT01-{base_name}-{counter}.docx"
    doc.save(out_docx)
    return out_docx

class App:
    def __init__(self, root):
        self.root = root
        root.title("CT01 Filler - Nhập thông tin công dân (offline)")
        root.geometry("1200x800")
        
        main_frame = tk.Frame(root, padx=15, pady=15)
        main_frame.pack(fill=tk.BOTH, expand=True)
        main_frame.rowconfigure(4, weight=1)
        main_frame.columnconfigure(0, weight=1)

        self.responsive_widgets = []
        self.member_fonts = []
        self.entries = {}

        info_frame = tk.Frame(main_frame)
        info_frame.grid(row=0, column=0, sticky="ew")
        info_frame.columnconfigure(1, weight=1)

        owner_question_frame = tk.Frame(main_frame)
        owner_question_frame.grid(row=1, column=0, sticky="ew", pady=5)
        owner_question_frame.columnconfigure(1, weight=1)

        owner_info_frame = tk.Frame(main_frame)
        owner_info_frame.grid(row=2, column=0, sticky="ew")
        owner_info_frame.columnconfigure(1, weight=1)

        main_fields = [
            ("Họ và tên:", "name"), ("Ngày sinh (dd/mm/yyyy):", "dob"),
            ("Giới tính:", "gender"), ("Số định danh (CCCD):", "id"),
            ("SĐT:", "phone")
        ]
        for i, (label_text, key) in enumerate(main_fields):
            label = tk.Label(info_frame, text=label_text, anchor="w")
            label.grid(row=i, column=0, sticky="w", padx=5, pady=5)
            entry = ttk.Combobox(info_frame, values=["Nam", "Nữ", "Khác"]) if key == "gender" else tk.Entry(info_frame)
            entry.grid(row=i, column=1, sticky="ew", padx=5, pady=5)
            self.entries[key] = entry
            self.responsive_widgets.extend([label, entry])

        is_owner_label = tk.Label(owner_question_frame, text="Bạn có phải chủ hộ không?", anchor="w")
        is_owner_label.grid(row=0, column=0, sticky="w", padx=5)
        self.responsive_widgets.append(is_owner_label)

        radio_frame = tk.Frame(owner_question_frame)
        radio_frame.grid(row=0, column=1, sticky="w")

        self.is_owner_var = tk.StringVar(value="Không")

        radio_yes = ttk.Radiobutton(radio_frame, text="Có", variable=self.is_owner_var, value="Có", command=self.autofill_owner_info)
        radio_yes.pack(side=tk.LEFT, padx=5)

        radio_no = ttk.Radiobutton(radio_frame, text="Không", variable=self.is_owner_var, value="Không", command=self.autofill_owner_info)
        radio_no.pack(side=tk.LEFT, padx=5)

        owner_fields = [
            ("Họ và tên chủ hộ:", "owner_name"),
            ("Mối quan hệ với chủ hộ:", "owner_relation"),
            ("Số định danh chủ hộ:", "owner_id")
        ]
        for i, (label_text, key) in enumerate(owner_fields):
            label = tk.Label(owner_info_frame, text=label_text, anchor="w")
            label.grid(row=i, column=0, sticky="w", padx=5, pady=5)
            entry = tk.Entry(owner_info_frame)
            entry.grid(row=i, column=1, sticky="ew", padx=5, pady=5)
            self.entries[key] = entry
            self.responsive_widgets.extend([label, entry])

        self.entries["name"].bind("<FocusOut>", self.update_title)

        self.lbl_members_title = tk.Label(main_frame, text="Những người xác nhận (mục 11):", anchor="w")
        self.lbl_members_title.grid(row=3, column=0, sticky="w", padx=5, pady=10)
        self.responsive_widgets.append(self.lbl_members_title)

        self.members_frame = tk.Frame(main_frame)
        self.members_frame.grid(row=4, column=0, sticky="nsew")
        self.members_frame.columnconfigure(0, weight=1)
        self.members = []
        self.add_member_ui()

        button_frame = tk.Frame(main_frame)
        button_frame.grid(row=5, column=0, sticky="sew", pady=10)
        button_frame.columnconfigure(0, weight=1)
        btn_add = tk.Button(button_frame, text="Thêm người phụ", command=self.add_member_ui)
        btn_add.pack(side=tk.LEFT)
        
        btn_auto_submit = tk.Button(button_frame, text="Tự động nộp hồ sơ", command=self.auto_click_nop_ho_so)
        btn_auto_submit.pack(side=tk.RIGHT, padx=5)
        
        btn_print_and_open = tk.Button(button_frame, text="In và Mở Cổng DVC", command=self.print_and_open_link)
        btn_print_and_open.pack(side=tk.RIGHT, padx=5)
        btn_export = tk.Button(button_frame, text="Xuất CT01 (Chỉ lưu file)", command=self.export_doc)
        btn_export.pack(side=tk.RIGHT)

        self.responsive_widgets.extend([btn_add, btn_export, btn_print_and_open, radio_yes, radio_no, btn_auto_submit])
        self.root.bind("<Configure>", self.on_resize)
        self.on_resize(None)
        self.update_title(None)

    def auto_click_nop_ho_so(self):
        """
        Sử dụng undetected-chromedriver để tránh bị phát hiện.
        """
        url_thu_tuc = "https://dichvucong.bocongan.gov.vn/bocongan/bothutuc/tthc?matt=26497"
        
        messagebox.showinfo("Bắt đầu tự động", f"Chuẩn bị mở trình duyệt và truy cập vào:\n{url_thu_tuc}")

        # <<< THAY ĐỔI QUAN TRỌNG: KHỞI TẠO DRIVER BẰNG uc.Chrome() >>>
        driver = uc.Chrome(enable_cdp_events=True, use_subprocess=True)
        
        try:
            driver.get(url_thu_tuc)
            driver.maximize_window()

            wait = WebDriverWait(driver, 20)
            
            messagebox.showinfo("Đang thực hiện", "Đang tìm nút 'Nộp hồ sơ'...")

            submit_button_selector = "a[href*='ma-thu-tuc-public=26497']"
            
            nop_ho_so_button = wait.until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, submit_button_selector))
            )
            
            nop_ho_so_button.click()
            messagebox.showinfo("Thành công!", "Đã nhấp vào nút 'Nộp hồ sơ' thành công!\nVui lòng tiếp tục thao tác trên trình duyệt.")
            
        except TimeoutException:
            messagebox.showerror("Lỗi", "Không tìm thấy nút 'Nộp hồ sơ' trong thời gian chờ. Trang web có thể đã thay đổi hoặc tải chậm.")
            if 'driver' in locals():
                driver.quit()
        except Exception as e:
            messagebox.showerror("Lỗi không mong muốn", f"Đã có lỗi xảy ra: {e}")
            if 'driver' in locals():
                driver.quit()

    def autofill_owner_info(self):
        owner_name_entry = self.entries["owner_name"]
        owner_rel_entry = self.entries["owner_relation"]
        owner_id_entry = self.entries["owner_id"]
        if self.is_owner_var.get() == "Có":
            main_name = self.entries["name"].get()
            main_id = self.entries["id"].get()
            owner_name_entry.delete(0, tk.END); owner_name_entry.insert(0, main_name)
            owner_rel_entry.delete(0, tk.END); owner_rel_entry.insert(0, "Chủ hộ")
            owner_id_entry.delete(0, tk.END); owner_id_entry.insert(0, main_id)
            owner_name_entry.config(state="readonly"); owner_rel_entry.config(state="readonly"); owner_id_entry.config(state="readonly")
        else:
            owner_name_entry.config(state="normal"); owner_rel_entry.config(state="normal"); owner_id_entry.config(state="normal")
            owner_name_entry.delete(0, tk.END); owner_rel_entry.delete(0, tk.END); owner_id_entry.delete(0, tk.END)

    def update_title(self, event):
        name = self.entries["name"].get().strip().upper()
        if not name: self.lbl_members_title.config(text="Những người xác nhận (mục 11):")
        else: self.lbl_members_title.config(text=f"Những người muốn xác nhận cùng với - {name}")

    def on_resize(self, event):
        width = self.root.winfo_width()
        if width > 1600: new_font_size, member_font_size = 28, 24
        elif width > 1200: new_font_size, member_font_size = 22, 19
        elif width > 800: new_font_size, member_font_size = 16, 14
        else: new_font_size, member_font_size = 12, 10

        main_font = ("Arial", new_font_size)
        for widget in self.responsive_widgets:
            try: widget.config(font=main_font)
            except tk.TclError: pass

        member_font = ("Arial", member_font_size)
        for font_list in self.member_fonts:
            for widget in font_list:
                try: widget.config(font=member_font)
                except tk.TclError: pass

    def add_member_ui(self):
        idx = len(self.members) + 1
        f = tk.Frame(self.members_frame, padx=10, pady=10, relief="ridge", bd=2)
        f.pack(fill=tk.X, expand=True, pady=5)
        f.columnconfigure(1, weight=1)
        labels_text = [f"Người #{idx} - Họ tên:", "Ngày sinh:", "Giới tính:", "Số định danh:", "Mối quan hệ:"]
        entries, member_widget_list = [], []
        for i, text in enumerate(labels_text):
            lbl = tk.Label(f, text=text, anchor="w")
            lbl.grid(row=i, column=0, sticky="w", padx=5, pady=2)
            entry = ttk.Combobox(f, values=["Nam","Nữ","Khác"]) if text == "Giới tính:" else tk.Entry(f)
            entry.grid(row=i, column=1, sticky="ew", padx=5, pady=2)
            entries.append(entry)
            member_widget_list.extend([lbl, entry])
        btn_delete = tk.Button(f, text="Xóa", command=lambda fr=f: self.remove_member(fr))
        btn_delete.grid(row=0, column=2, rowspan=2, padx=10)
        member_widget_list.append(btn_delete)
        self.members.append((f, *entries))
        self.member_fonts.append(member_widget_list)
        self.on_resize(None)

    def remove_member(self, frame):
        for i, t in enumerate(self.members):
            if t[0] == frame:
                t[0].destroy()
                self.members.pop(i)
                self.member_fonts.pop(i)
                break

    def _get_data_from_ui(self):
        main_info = {key: widget.get().strip() for key, widget in self.entries.items()}
        members = []
        for member_tuple in self.members:
            _, e_name, e_dob, e_gender, e_id, e_rel = member_tuple
            if e_name.get().strip():
                members.append({"name": e_name.get().strip(), "dob": e_dob.get().strip(),
                                 "gender": e_gender.get().strip(), "id_number": e_id.get().strip(),
                                 "relation": e_rel.get().strip()})
        return main_info, members

    def export_doc(self):
        try:
            main_info, members = self._get_data_from_ui()
            if not main_info.get("name"):
                messagebox.showwarning("Thiếu thông tin", "Vui lòng nhập họ và tên của người khai.")
                return None
            out_docx = fill_template(main_info, members)
            messagebox.showinfo("Hoàn tất", f"Đã tạo file Word thành công:\n{out_docx}")
            return out_docx
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không xuất được file: {e}")
            return None

    def print_and_open_link(self):
        # --- BƯỚC 1: Xuất và gửi lệnh in file CT01 ---
        docx_path = self.export_doc()
        if docx_path and os.path.exists(docx_path):
            try:
                # Gửi lệnh in file Word vừa tạo
                os.startfile(docx_path, "print")
                messagebox.showinfo("Đã gửi lệnh in", f"Đã gửi lệnh in cho file:\n{docx_path}")
            except Exception as e:
                messagebox.showerror("Lỗi in", f"Không thể gửi lệnh in.\nLỗi: {e}")

        # --- BƯỚC 2: Mở trực tiếp trang nộp hồ sơ online ---
        url_thu_tuc = "https://dichvucong.bocongan.gov.vn/bocongan/bothutuc/tthc?matt=26497"
        
        messagebox.showinfo("Chuẩn bị mở Web", f"Bây giờ chương trình sẽ mở trình duyệt và truy cập thẳng vào trang nộp hồ sơ.\n\nURL: {url_thu_tuc}")

        try:
            # Dùng webbrowser để mở URL trong trình duyệt mặc định của người dùng
            webbrowser.open(url_thu_tuc)
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể mở liên kết web.\nLỗi: {e}")


if __name__ == "__main__":
    root = tk.Tk()
    app = App(root)
    root.mainloop()